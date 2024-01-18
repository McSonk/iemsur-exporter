# pip install python-docx
# pip install markdownify

import logging

import xml.etree.ElementTree as ET

from pickle import load
import base64
from markdownify import markdownify as md
from io import BytesIO

from iemsur.moodle.objects import Exam, Question, Option, File

from docx import Document

logging.getLogger('pypandoc').addHandler(logging.NullHandler())
logging.basicConfig(level=logging.DEBUG, format='%(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class ExamReader:
    def __init__(self):
        self.exam = None

    def __analyse_question__(self, question):
        q_obj = Question()
        q_obj.type =  question.attrib['type']
        num = question.find('name').find('text').text
        if num != '.' and num != '. ' and num != 'IM1':
            q_obj.number = int(num)

        q_text = question.find('questiontext')
        q_obj.desc = q_text.find('text').text

        q_file = q_text.find('file')

        if q_file is not None:
            q_obj.file = File()
            q_obj.file.name = q_file.attrib['name']
            q_obj.file.encoding = q_file.attrib['encoding']
            q_obj.file.content = q_file.text

        if q_obj.is_true_false():
            q_obj.tf_value = question.find('.//answer[@fraction="100"]').find('text').text == 'true'
        elif q_obj.is_multiple_choice():
            for answer in question.findall('answer'):
                opt = Option(answer.find('text').text)
                opt.correct = answer.attrib['fraction'] == '100'
                q_obj.add_option(opt)
        elif q_obj.is_drag_box():
            # Number of blank spaces: [[n]]
            blanks = q_obj.desc.count('[[')
            for i, option in enumerate(question.findall('dragbox')):
                opt = Option(option.find('text').text)
                # Just the first i elements are correct
                opt.correct = i < blanks
                q_obj.add_option(opt)
        elif q_obj.is_matching():
            for option in question.findall('subquestion'):
                opt = Option(option.find('text').text)
                opt.match_value = option.find('answer').find('text').text
                file = option.find('file')
                if file is not None:
                    opt.file = File()
                    opt.file.name = file.attrib['name']
                    opt.file.encoding = file.attrib['encoding']
                    opt.file.content = file.text
                q_obj.add_option(opt)
        else:
            logger.warning(f'Unknown type: {q_obj.type}')
            logger.debug(q_obj)
        self.exam.add_question(q_obj)

    def read(self, file_name):
        xml = ET.parse(file_name)
        quiz = xml.getroot()
        self.exam = Exam()

        for question in quiz:
            self.__analyse_question__(question)

class ExamWriter:
    def __init__(self, exam):
        self.exam = exam

    def write(self, file_name):
        document = Document()
        document.add_heading('Examen', 0)
        
        self.exam.sort()

        for question in self.exam.questions:
            document.add_paragraph(md(question.desc, strip=['p']), style='List Number')
            if question.file is not None:
                logger.debug('file detected!')
                img = base64.b64decode(question.file.content)
                file = BytesIO(img)
                document.add_picture(file)
            if question.is_true_false():
                document.add_paragraph('', style='List Bullet 2').add_run('Verdadero').bold = question.tf_value
                document.add_paragraph('', style='List Bullet 2').add_run('Falso').bold = not question.tf_value
            elif question.is_multiple_choice() or question.is_drag_box():
                for option in question.options:
                    document.add_paragraph('', style='List Bullet 2').add_run(md(option.option, strip=['p'])).bold = option.correct
            elif question.is_matching():
                table = document.add_table(rows=0, cols=2)
                for option in question.options:
                    cells = table.add_row().cells
                    cells[0].text = md(option.option, strip=['p'])
                    if option.file is not None:
                        logger.debug('file detected!')
                        img = base64.b64decode(option.file.content)
                        file = BytesIO(img)
                        cells[0].paragraphs[0].add_run().add_picture(file)
                    cells[1].text = option.match_value


        document.save(file_name)
